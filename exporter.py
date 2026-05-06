import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


PILAR_COLORS = {
    "Comercial": RGBColor(0xC0, 0x39, 0x2B),
    "Institucional": RGBColor(0x21, 0x6E, 0xAB),
    "Educativo": RGBColor(0x27, 0xAE, 0x60),
}


def _add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _pilar_color(pilar: str) -> RGBColor:
    for key, color in PILAR_COLORS.items():
        if key.lower() in pilar.lower():
            return color
    return RGBColor(0x55, 0x55, 0x55)


def export_to_word(items: list, client_name: str) -> bytes:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Cover
    cover_title = doc.add_heading("", level=0)
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover_title.runs[0] if cover_title.runs else cover_title.add_run()
    run.text = f"Conteúdos para Redes Sociais"
    run.font.size = Pt(26)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(client_name)
    r.font.size = Pt(18)
    r.font.bold = True

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f"Gerado em {datetime.now().strftime('%d/%m/%Y às %H:%M')}  •  {len(items)} conteúdos")
    info.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()

    # Summary table
    doc.add_heading("Sumário de Conteúdos", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    headers = table.rows[0].cells
    for i, h in enumerate(["#", "Tema", "Pilar", "Formato"]):
        headers[i].text = h
        headers[i].paragraphs[0].runs[0].font.bold = True

    for item in items:
        row = table.add_row().cells
        row[0].text = str(item["num"])
        row[1].text = item["tema"]
        row[2].text = item["pilar"]
        row[3].text = item["formato"]

    doc.add_page_break()

    # Content pages
    for item in items:
        heading = doc.add_heading(f"#{item['num']} — {item['tema']}", level=1)
        heading.runs[0].font.color.rgb = _pilar_color(item["pilar"])

        meta = doc.add_paragraph()
        r1 = meta.add_run("Pilar: ")
        r1.bold = True
        r1.font.color.rgb = _pilar_color(item["pilar"])
        meta.add_run(item["pilar"])
        meta.add_run("     ")
        r2 = meta.add_run("Formato: ")
        r2.bold = True
        meta.add_run(item["formato"])

        doc.add_paragraph()

        content_para = doc.add_paragraph(item["conteudo"])
        content_para.paragraph_format.space_after = Pt(12)

        _add_horizontal_rule(doc)
        doc.add_paragraph()

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
